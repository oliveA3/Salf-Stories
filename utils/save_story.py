import re
import os
from docx import Document
from docx.shared import Pt
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
import base64


def add_formatted_paragraph(doc, text):
    if not text.strip():
        return

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing = Pt(18)

    tokens = re.split(r'(_[^_]+_)', text)

    for token in tokens:
        run = paragraph.add_run()
        run.font.name = 'Arial'
        run.font.size = Pt(12)

        if token.startswith("_") and token.endswith("_"):
            run.italic = True
            run.text = token[1:-1]
        else:
            run.text = token


def generate_and_save_story(title, content, repo, update=False):
    doc = Document()

    title_paragraph = doc.add_paragraph()
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(12)
    title_paragraph.paragraph_format.line_spacing = Pt(18)

    for line in content.splitlines():
        add_formatted_paragraph(doc, line)

    file_stream = BytesIO()

    doc.save(file_stream)
    file_stream.seek(0)
    file_bytes = file_stream.read()

    file_path = f"stories/{title}.docx"
    project_root = os.path.abspath(
        os.path.join(os.path.dirname(__file__), ".."))
    local_path = os.path.join(project_root, file_path)

    os.makedirs(os.path.dirname(local_path), exist_ok=True)
    with open(local_path, "wb") as f:
        f.write(file_bytes)

    if update:
        existing_file = repo.get_contents(file_path)
        repo.update_file(
            path=file_path,
            message=f"Updated story {title}",
            content=file_bytes,
            sha=existing_file.sha
        )
    else:
        repo.create_file(
            path=file_path,
            message=f"New story {title}",
            content=file_bytes
        )
