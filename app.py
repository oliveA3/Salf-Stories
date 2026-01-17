from flask import Flask, request, send_from_directory, request, jsonify
from docx import Document
from io import BytesIO
import base64
from github import Github, Auth
import os
import json
from dotenv import load_dotenv

load_dotenv()

app = Flask(__name__, static_folder="frontend")

# GITHUB VARIABLES
auth = Auth.Token(os.environ["GITHUB_TOKEN"])
g = Github(auth=auth)
repo = g.get_repo("oliveA3/Salf-Stories")


@app.route("/")
def home():
    return send_from_directory(app.static_folder, "index.html")


@app.route("/<path:path>")
def static_files(path):
    return send_from_directory(app.static_folder, path)


# LOADING EXISTING STORIES
@app.route("/stories_list")
def stories_list():
    files = repo.get_contents("stories")
    all_files = [f.name for f in files if f.name.endswith(".docx")]

    try:
        order_file = repo.get_contents("order.json")
        order_json = order_file.decoded_content.decode("utf-8")
        ordered_files = json.loads(order_json)
        ordered_files = [f for f in ordered_files if f in all_files]
        unordered_files = [f for f in all_files if f not in ordered_files]
        final_order = ordered_files + unordered_files
    except:
        final_order = all_files

    stories = []
    for file in files:
        if file.name in final_order:
            stories.append({
                "name": file.name,
                "url": file.download_url
            })

    stories.sort(key=lambda x: final_order.index(x["name"]))
    return jsonify(stories)


# SAVING NEW ORDER
@app.route("/save_order", methods=["POST"])
def save_order():
    new_order = request.json.get("order", [])

    with open("order.json", "w") as f:
        json.dump(new_order, f, indent=2)

    return jsonify({"message": "Order updated"})


# UPLOAD NEW STORY AS A .DOCX FILE TO THE GITHUB REPOSITORY
@app.route("/new_story", methods=["POST"])
def new_story():
    title = request.form.get("title")
    content = request.form.get("content")

    if not title or not content:
        return jsonify({"error": "Missing title or content"}), 400

    # Crear documento .docx
    doc = Document()
    title_run = doc.add_paragraph().add_run(title)
    title_run.bold = True
    doc.add_paragraph()
    content_run = doc.add_paragraph().add_run(content)

    # Guardar en memoria
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    file_name = f"stories/{title}.docx"

    # Subir a GitHub como binario (sin codificar)
    repo.create_file(
        path=file_name,
        message=f"New story {title}",
        content=file_stream.read()  # ← contenido binario directamente
    )

    return jsonify({"message": f"New story '{title}' uploaded to {file_name}"})


if __name__ == "__main__":
    app.run(debug=True)
